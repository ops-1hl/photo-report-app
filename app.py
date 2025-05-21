import streamlit as st
import pandas as pd
from PIL import Image
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
import os
import datetime

st.set_page_config(layout="centered")
st.title("📸 Gerador de Relatório - Oleões 2025")

# Uploads
excel_file = st.file_uploader("📄 Carregar ficheiro Excel (.xlsx)", type=["xlsx"])
images = st.file_uploader("🖼️ Carregar fotos (.jpg ou .png)", type=["jpg", "jpeg", "png"], accept_multiple_files=True)
logo_file = st.file_uploader("🏢 Logotipo da empresa", type=["jpg", "jpeg", "png"])
ghg_logo_file = st.file_uploader("🌍 Logotipo de certificação GHG", type=["jpg", "jpeg", "png"])

if excel_file and images and logo_file and ghg_logo_file:
    df = pd.read_excel(excel_file)
    st.success("✔️ Ficheiros carregados com sucesso!")

    document = Document()

    # Landscape config
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    photo_map = {os.path.splitext(photo.name)[0]: photo for photo in images}

    # COVER PAGE
    document.add_paragraph().add_run()
    table = document.add_table(rows=1, cols=2)
    row = table.rows[0]
    row.cells[0].width = Cm(8)
    row.cells[1].width = Cm(20)

    # Logo (Left)
    logo_img = Image.open(logo_file).convert("RGB")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_logo:
        logo_img.save(tmp_logo.name, format="JPEG")
        row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row.cells[0].paragraphs[0].add_run().add_picture(tmp_logo.name, width=Cm(6))

    # Title (Right)
    p = row.cells[1].paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("Município de Lisboa\nRelatório Final de Instalações\n\nAlargamento Rede de Oleões 2025")
    run.bold = True
    run.font.size = Pt(20)

    # GHG logo at bottom
    table = document.add_table(rows=1, cols=2)
    row = table.rows[0]
    row.cells[0].width = Cm(18)
    row.cells[1].width = Cm(10)
    p = row.cells[0].paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.add_run("GHG savings certified by:")
    ghg_img = Image.open(ghg_logo_file).convert("RGB")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_ghg:
        ghg_img.save(tmp_ghg.name, format="JPEG")
        row.cells[1].paragraphs[0].add_run().add_picture(tmp_ghg.name, width=Cm(3))

    document.add_page_break()

    # MAIN CONTENT
    for _, row in df.iterrows():
        codigo = str(row["ID"])
        document.add_paragraph(f"📌 Código do Oleão: {codigo}").bold = True

        if codigo in photo_map:
            try:
                img = Image.open(photo_map[codigo]).convert("RGB")
                resized = img.resize((int(12 * 37.8), int(16 * 37.8)))  # Resize to 12cm x 16cm @ 96 DPI
                with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as img_tmp:
                    resized.save(img_tmp.name, format="JPEG")
                    document.add_picture(img_tmp.name, width=Cm(12), height=Cm(16))
            except Exception as e:
                st.error(f"Erro com a imagem '{codigo}': {e}")
                document.add_paragraph("❌ Imagem inválida.")
        else:
            document.add_paragraph("❌ Foto não encontrada.")

        document.add_page_break()

    # FINAL PAGE
    final_section = document.add_section(WD_ORIENT.LANDSCAPE)
    final_section.page_width, final_section.page_height = section.page_width, section.page_height
    final_section.left_margin = Cm(2)
    final_section.right_margin = Cm(2)
    final_section.top_margin = Cm(1.5)
    final_section.bottom_margin = Cm(1.5)

    table = document.add_table(rows=1, cols=2)
    row = table.rows[0]
    row.cells[0].width = Cm(8)
    row.cells[1].width = Cm(20)

    # Final logo left
    row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row.cells[0].paragraphs[0].add_run().add_picture(tmp_logo.name, width=Cm(6))

    # Final message right
    p = row.cells[1].paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("Relatório Final\nRede de Oleões 2025")
    run.bold = True
    run.font.size = Pt(20)
    p.add_run(f"\n\nEmitido em: {datetime.date.today().strftime('%d/%m/%Y')}")

    # GHG logo bottom
    table = document.add_table(rows=1, cols=2)
    row = table.rows[0]
    row.cells[0].width = Cm(18)
    row.cells[1].width = Cm(10)
    p = row.cells[0].paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.add_run("GHG savings certified by:")
    row.cells[1].paragraphs[0].add_run().add_picture(tmp_ghg.name, width=Cm(3))

    # DOCX Export
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        document.save(tmp.name)
        with open(tmp.name, "rb") as f:
            docx_data = f.read()
        st.download_button("⬇️ Download DOCX", docx_data, file_name="relatorio_oleoes.docx")

    # PDF Export (Only works locally)
    if st.button("💾 Exportar como PDF (local apenas)"):
        try:
            from docx2pdf import convert
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, "report.docx")
                pdf_path = os.path.join(tmpdir, "report.pdf")
                document.save(docx_path)
                convert(docx_path, pdf_path)
                with open(pdf_path, "rb") as pdf_file:
                    st.download_button("📥 Baixar PDF", pdf_file, file_name="relatorio_oleoes.pdf")
        except Exception as e:
            st.error(f"Erro ao exportar como PDF: {e}")
